"""Inspect SOP-014 .docx so we can target edits precisely."""
import sys
from docx import Document

PATH = r"C:\Users\jgran\Desktop\SOP-014_Master_Inventory_Movement.docx"

doc = Document(PATH)

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"[{i}] ({p.style.name}) {t[:160]}")

print("\n=== TABLES ===")
for ti, tbl in enumerate(doc.tables):
    print(f"\n--- Table {ti} | {len(tbl.rows)} rows x {len(tbl.columns)} cols ---")
    for ri, row in enumerate(tbl.rows):
        cells = [c.text.strip().replace("\n", " ")[:40] for c in row.cells]
        print(f"  r{ri}: {cells}")

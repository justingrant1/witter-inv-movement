"""
Sync SOP-014 with the app/Airtable change that adds an explicit "Arrive HQ"
movement type (Return -> Returned, then Arrive HQ -> At HQ).

- Makes a timestamped backup first (clean copy, no track-changes).
- Targeted paragraph/run edits only; layout and styles untouched.
"""
import shutil
from datetime import datetime
from docx import Document

PATH = r"C:\Users\jgran\Desktop\SOP-014_Master_Inventory_Movement.docx"
BACKUP = PATH.replace(
    ".docx", f"_backup_{datetime.now():%Y%m%d_%H%M%S}.docx"
)

shutil.copy2(PATH, BACKUP)
print(f"Backup written: {BACKUP}")

doc = Document(PATH)


def set_paragraph_text(p, new_text):
    """Replace a paragraph's text while keeping its first run's formatting."""
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)


edits = 0

for p in doc.paragraphs:
    t = p.text

    # 1) Appendix C — Movements field list: add "Arrive HQ" to the Type vocab.
    if t.startswith("Movements — the custody log.") and "Arrive HQ" not in t:
        new = t.replace(
            "Booking pull / Ship / Return)",
            "Booking pull / Ship / Return / Arrive HQ)",
        )
        set_paragraph_text(p, new)
        edits += 1
        print("Edited: Appendix C Movements field list")

    # 2) §6.9 Returns — make the two-step state transition explicit.
    elif t.startswith("Surprise-set leftovers are returned to Lombard HQ"):
        new = (
            "Surprise-set leftovers are returned to Lombard HQ for processing; "
            "auction leftovers go back into the auction safe (Movement type: "
            "Return \u2192 status Returned). When the item physically arrives "
            "back at Lombard it is logged with Movement type: Arrive HQ "
            "\u2192 status At HQ, which clears the returns watchdog."
        )
        set_paragraph_text(p, new)
        edits += 1
        print("Edited: \u00a76.9 Returns two-step transition")

# 3) Lifecycle table (Table 2): ensure "At HQ" row notes the Arrive HQ trigger.
for tbl in doc.tables:
    header = [c.text.strip() for c in tbl.rows[0].cells]
    if header[:2] == ["State", "Meaning"]:
        for row in tbl.rows:
            if row.cells[0].text.strip() == "At HQ":
                cell = row.cells[1]
                if "Arrive HQ" not in cell.text:
                    p0 = cell.paragraphs[0]
                    set_paragraph_text(
                        p0,
                        "Back at Lombard, reconciled (logged via Arrive HQ)",
                    )
                    edits += 1
                    print("Edited: lifecycle table At HQ row")

# 4) Appendix C automations note — single-entry now covers Arrive HQ too.
for p in doc.paragraphs:
    if p.text.startswith(
        "Logging a Movement auto-sets the game\u2019s Status and Current safe"
    ):
        if "all nine move types" not in p.text:
            set_paragraph_text(
                p,
                "Logging a Movement auto-sets the game\u2019s Status and "
                "Current safe from the move type (single-entry). This now "
                "covers all nine move types, including Arrive HQ \u2192 At HQ.",
            )
            edits += 1
            print("Edited: Appendix C automations note")

doc.save(PATH)
print(f"\nSaved {PATH} with {edits} edit(s).")
